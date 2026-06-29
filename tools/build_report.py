from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "BAO_CAO_SOCKET_DEMO.md"
OUTPUT = ROOT / "docs" / "BAO_CAO_SOCKET_DEMO.docx"

NAVY = "17365D"
BLUE = "2F5597"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
WHITE = "FFFFFF"
BLACK = "000000"

# narrative_proposal preset with named academic_vietnam overrides.
TOKENS = {
    "page_width": 8.5,
    "page_height": 11.0,
    "margin": 1.0,
    "header_distance": 0.492,
    "footer_distance": 0.492,
    "content_dxa": 9360,
    "table_indent_dxa": 120,
    "font": "Times New Roman",
    "body_size": 12,
    "body_after": 6,
    "body_line": 1.5,
    "h1_size": 16,
    "h1_before": 18,
    "h1_after": 10,
    "h2_size": 13.5,
    "h2_before": 12,
    "h2_after": 6,
    "h3_size": 12.5,
    "h3_before": 8,
    "h3_after": 4,
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TOKENS["table_indent_dxa"]))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)


def allocate_widths(rows: list[list[str]]) -> list[int]:
    cols = len(rows[0])
    scores = []
    for idx in range(cols):
        max_len = max(len(row[idx]) if idx < len(row) else 0 for row in rows)
        scores.append(max(6, min(max_len, 45)))
    if cols >= 5:
        scores = [max(5, min(s, 24)) for s in scores]
    total = sum(scores)
    widths = [max(720, round(TOKENS["content_dxa"] * s / total)) for s in scores]
    delta = TOKENS["content_dxa"] - sum(widths)
    widths[-1] += delta
    if widths[-1] < 600:
        shortage = 600 - widths[-1]
        widths[-1] = 600
        donor = max(range(len(widths) - 1), key=lambda i: widths[i])
        widths[donor] -= shortage
    return widths


def set_run_font(run, name: str, size: float, color: str = BLACK, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, instruction: str) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instruction)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = ""
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)


def create_decimal_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "280")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    level.extend([start, num_fmt, lvl_text, suff, p_pr])
    abstract.append(level)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(list(numbering).index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.194)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(TOKENS["page_width"])
    section.page_height = Inches(TOKENS["page_height"])
    section.top_margin = Inches(TOKENS["margin"])
    section.bottom_margin = Inches(TOKENS["margin"])
    section.left_margin = Inches(TOKENS["margin"])
    section.right_margin = Inches(TOKENS["margin"])
    section.header_distance = Inches(TOKENS["header_distance"])
    section.footer_distance = Inches(TOKENS["footer_distance"])
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = TOKENS["font"]
    normal._element.rPr.rFonts.set(qn("w:ascii"), TOKENS["font"])
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), TOKENS["font"])
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), TOKENS["font"])
    normal.font.size = Pt(TOKENS["body_size"])
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(TOKENS["body_after"])
    normal.paragraph_format.line_spacing = TOKENS["body_line"]

    heading_values = {
        "Heading 1": (TOKENS["h1_size"], TOKENS["h1_before"], TOKENS["h1_after"], NAVY),
        "Heading 2": (TOKENS["h2_size"], TOKENS["h2_before"], TOKENS["h2_after"], BLUE),
        "Heading 3": (TOKENS["h3_size"], TOKENS["h3_before"], TOKENS["h3_after"], NAVY),
    }
    for style_name, (size, before, after, color) in heading_values.items():
        style = styles[style_name]
        style.font.name = TOKENS["font"]
        style._element.rPr.rFonts.set(qn("w:ascii"), TOKENS["font"])
        style._element.rPr.rFonts.set(qn("w:hAnsi"), TOKENS["font"])
        style._element.rPr.rFonts.set(qn("w:eastAsia"), TOKENS["font"])
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = TOKENS["font"]
        style.font.size = Pt(TOKENS["body_size"])
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run("BÁO CÁO ĐỒ ÁN LẬP TRÌNH SOCKET - SOCKETDEMO")
    set_run_font(hr, "Arial", 8.5, MID_GRAY, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Trang ")
    set_run_font(fr, "Arial", 9, MID_GRAY)
    add_field(fp, "PAGE")


def add_inline_markup(paragraph, text: str, size: float | None = None) -> None:
    # Supports Markdown bold and inline code while preserving Vietnamese text.
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, TOKENS["font"], size or TOKENS["body_size"])
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, "Consolas", (size or TOKENS["body_size"]) - 1, NAVY)
        else:
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, TOKENS["font"], size or TOKENS["body_size"], bold=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, TOKENS["font"], size or TOKENS["body_size"])


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline_markup(p, text)


def add_code_block(doc: Document, lines: list[str], language: str) -> None:
    if language == "mermaid":
        label = doc.add_paragraph()
        label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label.paragraph_format.space_before = Pt(6)
        label.paragraph_format.space_after = Pt(4)
        run = label.add_run("Sơ đồ Mermaid (có thể chỉnh sửa)")
        set_run_font(run, "Arial", 9, BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.05
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F7FA")
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:color"), "D0D5DD")
        edge.set(qn("w:space"), "5")
        borders.append(edge)
    p_pr.append(borders)
    run = p.add_run("\n".join(lines))
    set_run_font(run, "Consolas", 8.3 if language == "mermaid" else 8.8, BLACK)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = len(rows[0])
    rows = [row + [""] * (cols - len(row)) for row in rows]
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    widths = allocate_widths(rows)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    font_size = 8.2 if cols >= 5 else 9.2 if cols >= 3 else 10

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, "FAFBFC")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 and cols >= 3 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            add_inline_markup(p, value, font_size)
            for run in p.runs:
                if r_idx == 0:
                    run.bold = True

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        raw = lines[idx].strip().strip("|")
        cells = [cell.strip() for cell in raw.split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            rows.append(cells)
        idx += 1
    return rows, idx


def build_document() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)

    idx = 0
    first_page = True
    in_code = False
    code_language = ""
    code_lines: list[str] = []
    skip_static_toc = False
    pending_page_break = False
    previous_was_numbered = False
    current_num_id: int | None = None

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        if in_code:
            if stripped.startswith("```"):
                add_code_block(doc, code_lines, code_language)
                in_code = False
                code_language = ""
                code_lines = []
            else:
                code_lines.append(line)
            idx += 1
            continue

        if stripped.startswith("```"):
            in_code = True
            code_language = stripped[3:].strip().lower()
            idx += 1
            continue

        if stripped == "\\newpage":
            pending_page_break = True
            first_page = False
            previous_was_numbered = False
            current_num_id = None
            idx += 1
            continue

        if stripped.startswith("|") and idx + 1 < len(lines) and lines[idx + 1].strip().startswith("|"):
            rows, idx = parse_table(lines, idx)
            add_table(doc, rows)
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if first_page and level == 1:
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_after = Pt(54)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(text)
                set_run_font(run, "Arial", 24, NAVY, bold=True)
            elif first_page and level == 2:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(28)
                run = p.add_run(text)
                set_run_font(run, TOKENS["font"], 15, BLUE, bold=True)
            else:
                p = doc.add_heading(text, level=level)
                if pending_page_break:
                    p.paragraph_format.page_break_before = True
                    pending_page_break = False
                if text.startswith("PHỤ LỤC ") and text not in {"PHỤ LỤC A. HƯỚNG DẪN BUILD VÀ CHẠY"}:
                    p.runs[0].font.size = Pt(14)
                    p.paragraph_format.space_before = Pt(8)
                    p.paragraph_format.space_after = Pt(6)
                if level == 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if text == "MỤC LỤC":
                    toc = doc.add_paragraph()
                    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
                    skip_static_toc = True
            previous_was_numbered = False
            current_num_id = None
            idx += 1
            continue

        if skip_static_toc:
            if stripped.startswith("*Số trang"):
                skip_static_toc = False
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(stripped.strip("*"))
                set_run_font(run, TOKENS["font"], 10, MID_GRAY, italic=True)
            idx += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if numbered:
            if pending_page_break:
                doc.add_page_break()
                pending_page_break = False
            if not previous_was_numbered or current_num_id is None:
                current_num_id = create_decimal_numbering(doc)
            p = doc.add_paragraph()
            apply_numbering(p, current_num_id)
            add_inline_markup(p, numbered.group(1))
            previous_was_numbered = True
            idx += 1
            continue
        if bullet:
            if pending_page_break:
                doc.add_page_break()
                pending_page_break = False
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markup(p, bullet.group(1))
            previous_was_numbered = False
            current_num_id = None
            idx += 1
            continue

        if not stripped:
            idx += 1
            continue

        if pending_page_break:
            doc.add_page_break()
            pending_page_break = False

        if first_page and stripped.startswith("**"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(5)
            add_inline_markup(p, stripped, 11.5)
        elif stripped.startswith("*") and stripped.endswith("*"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped.strip("*"))
            set_run_font(run, TOKENS["font"], 10, MID_GRAY, italic=True)
        else:
            add_body_paragraph(doc, stripped)
        previous_was_numbered = False
        current_num_id = None
        idx += 1

    doc.core_properties.title = "Báo cáo đồ án lập trình Socket - SocketDemo"
    doc.core_properties.subject = "C++17, Winsock2, TCP và HTTP/1.1 rút gọn"
    doc.core_properties.author = "[CẦN BỔ SUNG]"
    doc.core_properties.keywords = "Socket, Winsock2, TCP, HTTP, Client-Server, C++"
    doc.core_properties.comments = "Tạo từ nội dung đã kiểm chứng ngày 29/06/2026."
    doc.settings.element.append(OxmlElement("w:updateFields"))
    doc.settings.element[-1].set(qn("w:val"), "true")
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    try:
        build_document()
    except Exception as exc:
        print(f"Failed to build report: {exc}", file=sys.stderr)
        raise
